#!/usr/bin/env python3
"""Convertit un .md (sous-ensemble) en .docx propre via python-docx.
Gère : # ## ### titres, > citations, - listes à puces, 1. listes numérotées,
tableaux | a | b |, **gras** inline, `code` inline, --- séparateur.
Usage: python3 _md_to_docx.py fichier.md
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EMBER = RGBColor(0xB8, 0x52, 0x1C)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_inline(paragraph, text):
    """Ajoute du texte avec gras (**...**) et code (`...`)."""
    parts = re.split(r'(\*\*.+?\*\*|`.+?`|\*[^*\n]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.color.rgb = EMBER
        elif len(part) > 2 and part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def _absorb_continuation(lines, i, buf):
    """Absorbe les lignes de continuation indentées d'un item de liste
    (markdown : une puce peut être enroulée sur plusieurs lignes indentées)."""
    while i < len(lines):
        cur = lines[i].rstrip()
        if (not cur.strip() or not cur.startswith(' ')
                or re.match(r'^\s*[-*] ', cur) or re.match(r'^\s*\d+\. ', cur)):
            break
        buf.append(cur.strip()); i += 1
    return i


def convert(md_path):
    docx_path = md_path.rsplit('.', 1)[0] + '.docx'
    lines = open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Tableau
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|?$', lines[i + 1].strip()):
            header = [c.strip() for c in line.strip('|').split('|')]
            i += 2  # saute l'entête + le séparateur
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Light Grid Accent 2'
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].clear()
                add_inline(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[:len(header)]):
                    cells[j].paragraphs[0].clear()
                    add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        if not line.strip():
            i += 1
            continue
        # Bloc de code clôturé ``` ... ``` → rendu monospace encadré (sans les ```).
        if line.lstrip().startswith('```'):
            i += 1
            code = []
            while i < len(lines) and not lines[i].lstrip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1  # saute le ``` de fermeture
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            r = p.add_run('\n'.join(code))
            r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
        if line.strip() == '---':
            doc.add_paragraph('─' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3); i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2); i += 1
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1); i += 1
        elif line.startswith('> '):
            # Regroupe les lignes de citation consécutives en un paragraphe.
            buf = []
            while i < len(lines) and lines[i].rstrip().startswith('>'):
                buf.append(lines[i].rstrip().lstrip('>').strip()); i += 1
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_inline(p, ' '.join(buf))
            for r in p.runs:
                r.italic = True
                if not r.bold: r.font.color.rgb = GREY
        elif re.match(r'^\s*[-*] ', line):
            buf = [re.sub(r'^\s*[-*] ', '', line)]; i += 1
            i = _absorb_continuation(lines, i, buf)
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, ' '.join(buf))
        elif re.match(r'^\s*\d+\. ', line):
            buf = [re.sub(r'^\s*\d+\. ', '', line)]; i += 1
            i = _absorb_continuation(lines, i, buf)
            p = doc.add_paragraph(style='List Number'); add_inline(p, ' '.join(buf))
        else:
            # Consomme TOUJOURS la ligne courante (évite la boucle infinie sur une
            # ligne type '#hashtag' qui n'est pas un titre), puis regroupe les
            # lignes adjacentes en un paragraphe (corrige le gras sur 2 lignes).
            buf = [line.strip()]; i += 1
            while i < len(lines):
                cur = lines[i].rstrip()
                if (not cur.strip() or cur.startswith('#') or cur.startswith('>')
                        or cur.strip() == '---' or cur.startswith('|')
                        or cur.lstrip().startswith('```')
                        or re.match(r'^\s*[-*] ', cur) or re.match(r'^\s*\d+\. ', cur)):
                    break
                buf.append(cur.strip()); i += 1
            p = doc.add_paragraph()
            add_inline(p, ' '.join(buf))

    doc.save(docx_path)
    print('OK ->', docx_path)


if __name__ == '__main__':
    for path in sys.argv[1:]:
        convert(path)
